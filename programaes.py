import sys
import os
from PyQt5.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QLabel, QPushButton, QFileDialog, QMessageBox, QTextEdit
)
from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
from cryptography.hazmat.primitives import padding
from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
from cryptography.hazmat.backends import default_backend
from cryptography.hazmat.primitives import hashes
from base64 import urlsafe_b64encode, urlsafe_b64decode
import secrets
from docx import Document
from PyPDF2 import PdfReader
from openpyxl import load_workbook

class FileEncryptorApp(QWidget):
    def __init__(self):
        super().__init__()
        self.init_ui()

    def init_ui(self):
        self.setWindowTitle("File Encryption and Decryption with AES")
        self.setGeometry(200, 200, 600, 400)

        layout = QVBoxLayout()

        self.label = QLabel("Select a file to encrypt or decrypt:")
        layout.addWidget(self.label)

        self.select_button = QPushButton("Select File")
        self.select_button.clicked.connect(self.select_file)
        layout.addWidget(self.select_button)

        self.encrypt_button = QPushButton("Encrypt File")
        self.encrypt_button.clicked.connect(self.encrypt_file)
        self.encrypt_button.setEnabled(False)
        layout.addWidget(self.encrypt_button)

        self.decrypt_button = QPushButton("Decrypt File")
        self.decrypt_button.clicked.connect(self.decrypt_file)
        self.decrypt_button.setEnabled(False)
        layout.addWidget(self.decrypt_button)

        self.preview_label = QLabel("File Preview:")
        layout.addWidget(self.preview_label)

        self.text_preview = QTextEdit()
        self.text_preview.setReadOnly(True)
        layout.addWidget(self.text_preview)

        self.setLayout(layout)

        self.selected_file = None

    def select_file(self):
        options = QFileDialog.Options()
        options |= QFileDialog.ReadOnly
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Select a File", "", "All Files (*.*);;Encrypted Files (*.enc);;Word Files (*.docx);;Excel Files (*.xlsx);;PDF Files (*.pdf)", options=options
        )
        if file_path:
            self.selected_file = file_path
            self.label.setText(f"Selected file: {file_path}")
            self.update_buttons()
            self.preview_file()

    def update_buttons(self):
        if self.selected_file.endswith('.enc'):
            self.encrypt_button.setEnabled(False)
            self.decrypt_button.setEnabled(True)
        elif self.selected_file.endswith(('.docx', '.xlsx', '.pdf')):
            self.encrypt_button.setEnabled(True)
            self.decrypt_button.setEnabled(False)
        else:
            self.encrypt_button.setEnabled(False)
            self.decrypt_button.setEnabled(False)

    def preview_file(self):
        try:
            if self.selected_file.endswith('.enc'):
                with open(self.selected_file, 'rb') as file:
                    content = file.read(500).hex()  # Preview first 500 bytes as hex
                self.text_preview.setText(f"Encrypted File Preview (Hex):\n{content}")
            elif self.selected_file.endswith('.docx'):
                doc = Document(self.selected_file)
                content = "\n".join([p.text for p in doc.paragraphs])
                self.text_preview.setText(content)
            elif self.selected_file.endswith('.pdf'):
                reader = PdfReader(self.selected_file)
                content = "\n".join([page.extract_text() for page in reader.pages[:5]])  # Preview up to 5 pages
                self.text_preview.setText(content)
            elif self.selected_file.endswith('.xlsx'):
                wb = load_workbook(self.selected_file)
                sheet = wb.active
                content = "\n".join([
                    ",".join([str(cell.value) for cell in row]) for row in sheet.iter_rows(min_row=1, max_row=10)
                ])  # Preview first 10 rows
                self.text_preview.setText(content)
            elif self.selected_file.endswith('.txt'):
                with open(self.selected_file, 'r', encoding='utf-8', errors='ignore') as file:
                    content = file.read(1000)  # Preview first 1000 characters
                self.text_preview.setText(content)
            else:
                self.text_preview.setText("Preview not available for this file type.")
        except Exception as e:
            self.text_preview.setText(f"Error previewing file: {str(e)}")

    def encrypt_file(self):
        if not self.selected_file:
            QMessageBox.warning(self, "No File Selected", "Please select a file to encrypt.")
            return

        try:
            password = "securepassword"  # Replace with a securely generated password
            salt = secrets.token_bytes(16)
            key = self.derive_key(password.encode(), salt)

            with open(self.selected_file, 'rb') as f:
                file_data = f.read()

            encrypted_data, iv = self.aes_encrypt(file_data, key)

            encrypted_file_path = self.selected_file + ".enc"

            result = QMessageBox.question(
                self, "Save Encrypted File", f"Do you want to save the encrypted file as: {encrypted_file_path}?",
                QMessageBox.Yes | QMessageBox.No
            )

            if result == QMessageBox.Yes:
                with open(encrypted_file_path, 'wb') as f:
                    f.write(salt + iv + encrypted_data)
                self.selected_file = encrypted_file_path
                self.text_preview.setText(f"Encrypted Data (Hex):\n{(salt + iv + encrypted_data).hex()[:1000]}...")
                QMessageBox.information(self, "Success", f"File encrypted successfully: {encrypted_file_path}")
            else:
                QMessageBox.information(self, "Cancelled", "Encryption operation was cancelled.")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"An error occurred: {str(e)}")

    def decrypt_file(self):
        if not self.selected_file:
            QMessageBox.warning(self, "No File Selected", "Please select a file to decrypt.")
            return

        try:
            password = "securepassword"  # Replace with the same password used for encryption

            with open(self.selected_file, 'rb') as f:
                file_data = f.read()

            salt = file_data[:16]
            iv = file_data[16:32]
            encrypted_data = file_data[32:]

            key = self.derive_key(password.encode(), salt)
            decrypted_data = self.aes_decrypt(encrypted_data, key, iv)

            decrypted_file_path = self.selected_file.replace(".enc", "_decrypted")

            result = QMessageBox.question(
                self, "Save Decrypted File", f"Do you want to save the decrypted file as: {decrypted_file_path}?",
                QMessageBox.Yes | QMessageBox.No
            )

            if result == QMessageBox.Yes:
                with open(decrypted_file_path, 'wb') as f:
                    f.write(decrypted_data)
                self.selected_file = decrypted_file_path
                self.preview_decrypted_file(decrypted_file_path)
                QMessageBox.information(self, "Success", f"File decrypted successfully: {decrypted_file_path}")
            else:
                QMessageBox.information(self, "Cancelled", "Decryption operation was cancelled.")

        except Exception as e:
            self.text_preview.setText(f"Error decrypting file: {str(e)}")

    def preview_decrypted_file(self, decrypted_file_path):
        try:
            if decrypted_file_path.endswith('.docx'):
                doc = Document(decrypted_file_path)
                content = "\n".join([p.text for p in doc.paragraphs])
                self.text_preview.setText(content)
            elif decrypted_file_path.endswith('.pdf'):
                reader = PdfReader(decrypted_file_path)
                content = "\n".join([page.extract_text() for page in reader.pages[:5]])  # Preview up to 5 pages
                self.text_preview.setText(content)
            elif decrypted_file_path.endswith('.xlsx'):
                wb = load_workbook(decrypted_file_path)
                sheet = wb.active
                content = "\n".join([
                    ",".join([str(cell.value) for cell in row]) for row in sheet.iter_rows(min_row=1, max_row=10)
                ])  # Preview first 10 rows
                self.text_preview.setText(content)
            elif decrypted_file_path.endswith('.txt'):
                with open(decrypted_file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    content = file.read(1000)  # Preview first 1000 characters
                self.text_preview.setText(content)
            else:
                self.text_preview.setText("Preview not available for this file type.")
        except Exception as e:
            self.text_preview.setText(f"Error previewing decrypted file: {str(e)}")

    # KRIPTOGRAFI KEY
    def derive_key(self, password, salt): 
        kdf = PBKDF2HMAC( #PBKDF2HMAC: Algoritma derivasi kunci yang aman terhadap serangan brute-force. 
            algorithm=hashes.SHA256(),  # Algoritma hashing
            length=32, # Panjang kunci (32 byte = 256 bit untuk AES)
            salt=salt, #Salt: 16 byte data acak yang ditambahkan untuk membuat setiap kunci unik meskipun password sama.
            iterations=100000, #Iterasi ditetapkan ke 100.000 untuk meningkatkan keamanan.
            backend=default_backend()
        )
        return kdf.derive(password)

    def aes_encrypt(self, data, key):
        iv = secrets.token_bytes(16)
        cipher = Cipher(algorithms.AES(key), modes.CBC(iv), backend=default_backend())
        encryptor = cipher.encryptor()

        padder = padding.PKCS7(algorithms.AES.block_size).padder()
        padded_data = padder.update(data) + padder.finalize()

        encrypted_data = encryptor.update(padded_data) + encryptor.finalize()
        return encrypted_data, iv

    def aes_decrypt(self, encrypted_data, key, iv):
        cipher = Cipher(algorithms.AES(key), modes.CBC(iv), backend=default_backend())
        decryptor = cipher.decryptor()

        decrypted_padded_data = decryptor.update(encrypted_data) + decryptor.finalize()

        unpadder = padding.PKCS7(algorithms.AES.block_size).unpadder()
        data = unpadder.update(decrypted_padded_data) + unpadder.finalize()
        return data

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = FileEncryptorApp()
    window.show()
    sys.exit(app.exec_())
